#!/usr/bin/env python3
import sys, json, re
from collections import OrderedDict
from pathlib import Path
from docx import Document  # python-docx

# ---------- helpers ----------

def _is_bilingual_header(txt: str) -> bool:
    # e.g., "Project – Proje" or "Data Scope - Veri Kapsamı"
    return bool(re.match(r"^([^–\-]+?)\s*[–-]\s*([^–\-]+?)$", txt))

def _split_header(txt: str) -> str:
    if "–" in txt:
        return txt.split("–", 1)[0].strip()
    if " - " in txt:
        return txt.split(" - ", 1)[0].strip()
    return txt.strip()

def _is_list_paragraph(p) -> bool:
    """True if paragraph is a numbered/bulleted item."""
    try:
        pPr = p._p.pPr  # paragraph properties
        return (pPr is not None) and (pPr.numPr is not None)
    except AttributeError:
        return False

# ---------- main parser ----------

def parse_form(docx_path: Path) -> dict:
    doc = Document(str(docx_path))

    # Collect paragraph text + basic flags
    paras = []
    for p in doc.paragraphs:
        txt = (p.text or "").strip()
        style_name = (getattr(getattr(p, "style", None), "name", "") or "").lower()
        is_list = _is_list_paragraph(p)
        paras.append({"text": txt, "style": style_name, "is_list": is_list})

    # Header indices: bilingual only (avoids capturing values like personal names)
    header_idxs = [
        i for i, d in enumerate(paras)
        if d["text"] and _is_bilingual_header(d["text"]) and d["style"] != "title"
    ]

    result = OrderedDict()

    for j, start in enumerate(header_idxs):
        raw_key = paras[start]["text"]
        key = _split_header(raw_key)

        end = header_idxs[j + 1] if j + 1 < len(header_idxs) else len(paras)
        # section body = non-empty lines until next header
        body = [paras[k] for k in range(start + 1, end) if paras[k]["text"]]

        # if first body line itself looks like a header (rare formatting noise), skip it
        if body and _is_bilingual_header(body[0]["text"]):
            body = body[1:]

        # simple scalar value (single non-list line without "Label:")
        if body and len(body) == 1 and not body[0]["is_list"] and ":" not in body[0]["text"]:
            result[key] = body[0]["text"]
            continue

        # More complex: labeled subfields and/or narrative + list items
        sub = OrderedDict()
        narrative = []
        pending_label = None

        for row in body:
            t = row["text"]

            # label lines like "Sponsor:" or "Stakeholders: names"
            is_label = bool(re.match(r"^[^:]{2,}:\s*$", t)) or bool(re.match(r"^[^:]{2,}:\s+.+$", t))
            if is_label:
                # flush previous empty label if started
                if pending_label is not None and pending_label not in sub:
                    sub[pending_label] = ""
                label, _, after = t.partition(":")
                pending_label = label.strip()
                after = after.strip()
                if after:
                    sub[pending_label] = after
                else:
                    sub.setdefault(pending_label, "")
                continue

            # not a label line
            if pending_label is not None:
                existing = sub.get(pending_label, "")
                joiner = "\n" if existing else ""
                sub[pending_label] = existing + joiner + t
            else:
                narrative.append(("-ITEM-" if row["is_list"] else "TEXT", t))

        # normalize shape
        if sub:
            # convert multi-line values to list
            for k2, v2 in list(sub.items()):
                lines = [ln.strip() for ln in v2.splitlines() if ln.strip()]
                sub[k2] = lines if len(lines) > 1 else (lines[0] if lines else "")
            result[key] = sub
        else:
            items = [t for tag, t in narrative if tag == "-ITEM-"]
            texts = [t for tag, t in narrative if tag == "TEXT"]
            if items and texts:
                result[key] = {"overview": "\n\n".join(texts).strip(), "items": items}
            elif items:
                result[key] = items
            else:
                result[key] = "\n\n".join(texts).strip()

    # Also capture any tables as raw rows (optional, handy for scoring tables etc.)
    tables_data = []
    for tbl in doc.tables:
        rows = []
        for r in tbl.rows:
            cells = [c.text.strip() for c in r.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            tables_data.append(rows)

    if tables_data:
        result["_tables"] = tables_data

    return result

def main():
    in_doc = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("input.docx")
    out_json = Path(sys.argv[2]) if len(sys.argv) > 2 else Path("output.json")
    data = parse_form(in_doc)
    out_json.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out_json}")

if __name__ == "__main__":
    main()
